from pathlib import Path
from typing import List

import base64
import cv2
import docx
import numpy as np
import pdfplumber
import pytesseract
from pdf2image import convert_from_path

from .config import Settings


class ImagePreprocessor:
    def __init__(self, settings: Settings):
        self.settings = settings

    def load_cv(self, file_path: str) -> List[np.ndarray]:
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext == ".pdf":
            return self._load_pdf(file_path)
        if ext == ".docx":
            return self._load_docx(file_path)
        if ext in [".png", ".jpg", ".jpeg"]:
            img = cv2.imread(file_path)
            return [img] if img is not None else []
        raise ValueError(f"Format non supporté : {ext}")

    def _load_pdf(self, pdf_path: str) -> List[np.ndarray]:
        kwargs = {
            "dpi": self.settings.cv_dpi,
            "fmt": "jpeg",
        }
        if self.settings.poppler_path:
            kwargs["poppler_path"] = self.settings.poppler_path

        pages_pil = convert_from_path(pdf_path, **kwargs)

        images = []
        for pil_img in pages_pil:
            img_np = np.array(pil_img)
            img_bgr = cv2.cvtColor(img_np, cv2.COLOR_RGB2BGR)
            images.append(img_bgr)
        return images

    def _load_docx(self, docx_path: str) -> List[np.ndarray]:
        doc = docx.Document(docx_path)
        lines = []

        for para in doc.paragraphs:
            txt = para.text.strip()
            if txt:
                lines.append(txt)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    lines.append(row_text)

        if not lines:
            return [np.ones((800, 600, 3), dtype=np.uint8) * 255]

        font = cv2.FONT_HERSHEY_SIMPLEX
        font_scale = 0.38
        thickness = 1
        line_h = 18
        margin_x = 25
        margin_y = 30
        width = 900

        total_lines = len(lines)
        height = max(800, margin_y * 2 + total_lines * line_h + 50)
        height = min(height, 6000)

        img = np.ones((height, width, 3), dtype=np.uint8) * 255
        y = margin_y

        for line in lines:
            if y > height - margin_y:
                break

            max_chars = 120
            if len(line) > max_chars:
                segments = [line[i:i + max_chars] for i in range(0, len(line), max_chars)]
                for seg in segments:
                    if y > height - margin_y:
                        break
                    cv2.putText(
                        img, seg, (margin_x, y), font, font_scale,
                        (0, 0, 0), thickness, cv2.LINE_AA
                    )
                    y += line_h
            else:
                cv2.putText(
                    img, line, (margin_x, y), font, font_scale,
                    (0, 0, 0), thickness, cv2.LINE_AA
                )
                y += line_h

        return [img]

    def preprocess(self, img: np.ndarray) -> np.ndarray:
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        if np.mean(gray) < 100:
            lab = cv2.cvtColor(img, cv2.COLOR_BGR2LAB)
            l, a, b = cv2.split(lab)
            clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
            img = cv2.merge([clahe.apply(l), a, b])
            img = cv2.cvtColor(img, cv2.COLOR_LAB2BGR)
        return img

    def to_base64(self, img: np.ndarray) -> str:
        _, buf = cv2.imencode(".jpg", img, [cv2.IMWRITE_JPEG_QUALITY, 95])
        return base64.b64encode(buf).decode("utf-8")

    def extract_docx_text(self, docx_path: str) -> str:
        doc = docx.Document(docx_path)
        parts = []

        for para in doc.paragraphs:
            txt = para.text.strip()
            if txt:
                parts.append(txt)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)

    def extract_pdf_text(self, pdf_path: str) -> str:
        parts = []

        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    txt = (page.extract_text() or "").strip()
                    if txt:
                        parts.append(txt)
        except Exception:
            pass

        if parts:
            return "\n\n".join(parts)

        # Fallback OCR si aucun texte extractible
        try:
            pages = self._load_pdf(pdf_path)
            ocr_parts = []
            for img in pages:
                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                txt = pytesseract.image_to_string(gray)
                txt = txt.strip()
                if txt:
                    ocr_parts.append(txt)
            return "\n\n".join(ocr_parts)
        except Exception:
            return ""

    def extract_pdf_pages_text(self, pdf_path: str) -> List[str]:
        pages_text = []
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    txt = (page.extract_text() or "").strip()
                    pages_text.append(txt)
        except Exception:
            pages_text = []

        if any(pages_text):
            return pages_text

        try:
            pages = self._load_pdf(pdf_path)
            ocr_pages = []
            for img in pages:
                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                txt = pytesseract.image_to_string(gray)
                ocr_pages.append(txt.strip())
            return ocr_pages
        except Exception:
            return pages_text

        return pages_text

    def extract_text(self, file_path: str) -> str:
        path = Path(file_path)
        ext = path.suffix.lower()
        if ext == ".pdf":
            return self.extract_pdf_text(file_path)
        if ext == ".docx":
            return self.extract_docx_text(file_path)
        return ""
